import docx
import os

APP_ROOT = os.path.dirname(os.path.abspath(__file__))
OUTPUTS_FOLDER = os.path.join(APP_ROOT, 'outputs','generated_docx_files')

def write_docx_file(predicted_results):
    mydoc = docx.Document()
    mydoc.add_heading("Functional Requirements", 1)
    for key, value in predicted_results.items():
        if value == 'F':
            mydoc.add_paragraph(key)
    mydoc.add_paragraph("-----------------------------------------------------------------------------------------------------------------")
    mydoc.add_heading("Non Functional Requirements", 1)
    for key, value in predicted_results.items():
        if value != 'F':
            mydoc.add_paragraph(key)
    mydoc.save(OUTPUTS_FOLDER+"/Requirement.docx")
    print("Document Generated !")
    return True











